"""docx extraction helpers.

Two extraction modes:

* `extract_text(filepath)` — the user-specified helper: joins paragraph texts only.
  Used by the prose parsers (contract clauses, witness, expert, letter, record, pleading),
  where embedded tables (SOW milestone table, signature blocks) should be skipped.

* `iter_blocks(filepath)` — walks the document body in true document order, yielding both
  paragraphs and tables. Required by the table-aware parsers (email, internal_email,
  defect_log, and the Order Form path of the contract parser), because in the real .docx
  files the email headers, order-form line items and defect rows live inside Word tables,
  not paragraphs. A paragraph-only read would silently drop them.
"""

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def extract_text(filepath):
    """Join all paragraph texts (user-specified). Tables are not included."""
    doc = Document(filepath)
    return "\n".join(para.text for para in doc.paragraphs)


def iter_blocks(filepath):
    """Yield blocks in document order.

    Yields ('para', text) for each paragraph and
    ('table', rows) for each table, where rows is a list of lists of stripped cell texts.
    """
    doc = Document(filepath)
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            yield ("para", Paragraph(child, doc).text)
        elif tag.endswith("}tbl"):
            tbl = Table(child, doc)
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            yield ("table", rows)
